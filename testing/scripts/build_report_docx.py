#!/usr/bin/env python3
"""Generator LAPORAN PENGUJIAN (Word .docx) UNISMUH ComputeHub.

Menyusun laporan QA end-to-end lengkap dengan SCREENSHOT (dipilih otomatis: versi
TERBARU per tampilan dari folder testing/screenshots/). Dijalankan dari folder testing:

    ../backend/.venv/bin/python scripts/build_report_docx.py

Output: testing/Laporan-Pengujian-UNISMUH-ComputeHub.docx
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

TESTING = Path(__file__).resolve().parent.parent
SHOTS = TESTING / "screenshots"
OUT = TESTING / "Laporan-Pengujian-UNISMUH-ComputeHub.docx"

_TMP = Path(tempfile.mkdtemp(prefix="chreport-"))
_prep: dict[tuple[str, int], Path] = {}


def prepared(src: Path, max_w: int) -> Path:
    """Kecilkan gambar (resize + JPEG q85) agar .docx ringan & mudah dibagikan."""
    key = (str(src), max_w)
    if key in _prep:
        return _prep[key]
    img = Image.open(src).convert("RGB")
    if img.width > max_w:
        h = round(img.height * max_w / img.width)
        img = img.resize((max_w, h), Image.LANCZOS)
    out = _TMP / (src.stem + f"_{max_w}.jpg")
    img.save(out, "JPEG", quality=85, optimize=True)
    _prep[key] = out
    return out

BRAND = RGBColor(0x1D, 0x4E, 0xD8)   # biru brand
DARK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
OK = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xB9, 0x1C, 0x1C)


# --------------------------------------------------------------- util screenshot
def pick(rel_dir: str, suffix: str) -> Path | None:
    """Ambil screenshot TERBARU (timestamp terbesar) utk <rel_dir>/<ts>-<suffix>.png."""
    d = SHOTS / rel_dir
    if not d.is_dir():
        return None
    best_ts, best = -1, None
    for f in d.glob("*.png"):
        m = re.match(r"(\d+)-(.+)\.png$", f.name)
        if m and m.group(2) == suffix:
            ts = int(m.group(1))
            if ts > best_ts:
                best_ts, best = ts, f
    return best


# --------------------------------------------------------------- util dokumen
def set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tcPr.append(shd)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gambar: " + text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def add_shot(doc: Document, rel_dir: str, suffix: str, cap: str, width_cm: float = 16.0) -> bool:
    f = pick(rel_dir, suffix)
    if not f:
        return False
    max_w = 1200 if width_cm >= 12 else 760
    img = prepared(f, max_w)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Cm(width_cm))
    caption(doc, cap)
    return True


def add_shots_grid(doc: Document, items, width_cm: float) -> None:
    """Sematkan beberapa screenshot (list of (dir,suffix,cap))."""
    for rel_dir, suffix, cap in items:
        add_shot(doc, rel_dir, suffix, cap, width_cm)


def h(doc: Document, text: str, level: int = 1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.color.rgb = BRAND if level <= 2 else DARK
    return hd


def para(doc: Document, text: str, size: int = 10, color=None, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullets(doc: Document, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it).font.size = Pt(10)


def table(doc: Document, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, hh in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(hh)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1D4ED8")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    return t


# =============================================================== BUILD
doc = Document()
# margin & font dasar
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)

# ----- SAMPUL
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN PENGUJIAN PERANGKAT LUNAK")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = BRAND
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNISMUH ComputeHub")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = DARK
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Academic HPC Platform — Program Studi Informatika, Universitas Muhammadiyah Makassar")
r.font.size = Pt(11)
r.font.color.rgb = MUTED
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Pengujian Otomatis End-to-End · API · Keamanan · Performa · Responsif · Per-Peran")
r.font.size = Pt(11)
r.italic = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Kerangka uji: Playwright (Chromium) — 7 project pengujian")
r.font.size = Pt(11)
for _ in range(2):
    doc.add_paragraph()
# ringkas hasil di sampul
tb = table(
    doc,
    ["Metrik", "Nilai"],
    [
        ["Total kasus uji", "92"],
        ["Lulus", "90"],
        ["Skip (kondisional sah)", "2"],
        ["Gagal", "0"],
        ["Flaky", "0"],
        ["Durasi", "± 2,1 menit"],
        ["Tanggal", "8–9 Juli 2026"],
        ["Build", "main (DB PostgreSQL lokal)"],
    ],
    widths=[7, 8],
)
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
doc.add_paragraph()
para(
    doc,
    "0 kegagalan fungsional · 0 celah otorisasi/injeksi · 0 error fatal JavaScript.",
    size=11, color=OK, bold=True,
)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ----- DAFTAR ISI (statis)
h(doc, "Daftar Isi", 1)
toc = [
    "1. Ringkasan Eksekutif",
    "2. Konteks & Prinsip Pengujian",
    "3. Lingkungan & Alat Uji",
    "4. Metodologi & Cakupan Project",
    "5. Hasil Fungsional per Fitur (dengan tangkapan layar)",
    "6. Pengujian Per-Peran (Super Admin · Admin · Dosen · Mahasiswa)",
    "7. Pengujian Keamanan",
    "8. Pengujian Performa",
    "9. Pengujian Responsif (Mobile & Tablet)",
    "10. Temuan (Bug) & Perbaikan",
    "11. Observasi",
    "12. Inventaris Kasus Uji",
    "13. Kesimpulan",
    "Lampiran: Cara Menjalankan & Artefak",
]
for t in toc:
    para(doc, t, size=10)
doc.add_page_break()

# ----- 1. RINGKASAN EKSEKUTIF
h(doc, "1. Ringkasan Eksekutif", 1)
para(
    doc,
    "Dokumen ini melaporkan pengujian otomatis menyeluruh terhadap aplikasi UNISMUH "
    "ComputeHub — platform komputasi GPU (HPC) untuk sivitas akademika. Pengujian "
    "dilakukan memakai Playwright pada peramban Chromium melalui 7 project pengujian "
    "(publik, API, keamanan, desktop, mobile, tablet, performa) yang mencakup seluruh "
    "rute aplikasi, otorisasi keempat peran pengguna, keamanan (OWASP), performa, dan "
    "keterbacaan pada berbagai ukuran layar.",
)
para(
    doc,
    "Hasil akhir: dari 92 kasus uji, 90 LULUS, 2 di-skip secara sah (kondisional), "
    "0 GAGAL, dan 0 flaky, dengan durasi sekitar 2,1 menit. Dua temuan UX (BUG-001 & "
    "BUG-002) telah diperbaiki dan regresinya dikunci oleh kasus uji. Tidak ditemukan "
    "celah otorisasi, injeksi, maupun error fatal.",
)
h(doc, "Sorotan", 2)
bullets(doc, [
    "Fitur baru 'Unduh folder & seluruh workspace (.zip)' diuji penuh (API + UI) — semua LULUS.",
    "Cakupan 4 peran: Super Admin, Admin, Dosen, dan Mahasiswa (matriks otorisasi API + UI role-aware).",
    "Privilege-escalation (mahasiswa → endpoint admin) DIBUKTIKAN ditolak (HTTP 403).",
    "Keamanan: header anti-clickjacking, HSTS, CSP; directory traversal, SQLi, dan XSS ditolak.",
    "Bersifat NON-DESTRUKTIF: dijalankan pada server produksi bersama tanpa mengganggu pengguna nyata.",
])

# ----- 2. KONTEKS
h(doc, "2. Konteks & Prinsip Pengujian", 1)
para(
    doc,
    "Pengujian dijalankan pada server yang dipakai pengguna nyata (produksi bersama). "
    "Karena itu seluruh rangkaian uji dirancang NON-DESTRUKTIF:",
)
bullets(doc, [
    "Autentikasi via injeksi token (reuse session_token akun yang ada) — tidak menendang sesi pengguna.",
    "Operasi berkas hanya pada /persist milik akun UJI sendiri (terisolasi, reversibel).",
    "Tanpa load-test masif (risiko DoS) dan tanpa menguras rate-limit login (risiko mengunci pengguna).",
    "Keamanan diuji dengan memverifikasi aplikasi MENOLAK serangan — bukan mengeksploitasi/merusak.",
    "Server headless (tanpa layar) → mode headless + video + trace + screenshot memberi peninjauan visual setara.",
])

# ----- 3. LINGKUNGAN
h(doc, "3. Lingkungan & Alat Uji", 1)
table(
    doc,
    ["Komponen", "Detail"],
    [
        ["Aplikasi", "UNISMUH ComputeHub (React + Vite + FastAPI)"],
        ["Base URL uji", "http://127.0.0.1:8088 (same-origin dengan backend)"],
        ["Basis data", "PostgreSQL lokal (Docker) — latensi rendah"],
        ["Kerangka uji", "Playwright (Chromium bundel Playwright)"],
        ["Mode", "Headless · 2 workers · retry 1 · slowMo 200ms"],
        ["Artefak", "Screenshot · Video · Trace (per langkah) · HTML/JUnit/JSON report"],
        ["Server", "hpc-ai — 2× NVIDIA L40S 46 GB (produksi bersama)"],
    ],
    widths=[5, 11],
)

# ----- 4. METODOLOGI
h(doc, "4. Metodologi & Cakupan Project", 1)
para(doc, "Suite terbagi menjadi 7 project pengujian:")
table(
    doc,
    ["Project", "Peran / Viewport", "Lulus", "Skip", "Gagal"],
    [
        ["public", "tanpa auth (Desktop Chrome)", "7", "0", "0"],
        ["api", "bearer admin/super-admin/student/dosen", "15", "1", "0"],
        ["security", "context per-peran", "9", "0", "0"],
        ["desktop", "admin + mahasiswa/dosen · 1440×900", "42", "1", "0"],
        ["mobile", "Pixel 7 (412×839)", "5", "0", "0"],
        ["tablet", "820×1180", "5", "0", "0"],
        ["performance", "admin · 1440×900", "7", "0", "0"],
        ["TOTAL", "—", "90", "2", "0"],
    ],
    widths=[3, 6.5, 2, 2, 2],
)

# ----- 5. FUNGSIONAL PER FITUR
h(doc, "5. Hasil Fungsional per Fitur (dengan tangkapan layar)", 1)

h(doc, "5.1 Autentikasi & Halaman Publik", 2)
para(doc, "Login, validasi form, penolakan kredensial salah, toggle password, landing, "
          "halaman 404, dan pengalihan rute terproteksi ke /welcome.")
add_shots_grid(doc, [
    ("login", "landing", "Halaman Landing/Welcome (publik)"),
    ("login", "login-form", "Halaman Login"),
    ("login", "empty-validation", "Validasi saat form dikosongkan"),
    ("login", "after-invalid-login", "Kredensial salah ditolak (tanpa membocorkan info)"),
    ("login", "password-shown", "Toggle tampilkan kata sandi"),
    ("login", "protected-redirect", "Rute terproteksi dialihkan ke /welcome (belum login)"),
    ("login", "not-found", "Halaman 404 untuk rute tak dikenal"),
], 15)

h(doc, "5.2 Dashboard", 2)
para(doc, "Dashboard menampilkan ringkasan job, status GPU/CPU/RAM, dan kebijakan eksekusi (admin).")
add_shot(doc, "dashboard", "overview", "Dashboard (peran admin)")

h(doc, "5.3 Daftar Job", 2)
para(doc, "Daftar job, filter status, dan halaman detail job. Eksekusi job nyata sengaja "
          "tidak dipicu untuk melindungi antrian/GPU produksi.")
add_shots_grid(doc, [
    ("jobs", "list", "Daftar Job"),
    ("jobs", "filter-applied", "Filter status job diterapkan"),
    ("jobs", "detail", "Detail job"),
], 15)

h(doc, "5.4 Penyimpanan & Unduh Folder (fitur baru)", 2)
para(doc, "Halaman Penyimpanan (workspace persisten /persist) — indikator kuota, unggah "
          "berkas, dan FITUR BARU: unduh folder & seluruh workspace sebagai .zip "
          "(tombol 'Unduh semua').")
add_shots_grid(doc, [
    ("storage", "view", "Halaman Penyimpanan + indikator kuota"),
    ("storage", "upload-button", "Tombol Unggah"),
    ("storage", "after-upload", "Berkas QA muncul setelah diunggah"),
    ("storage", "download-all", "Tombol 'Unduh semua' — fitur unduh seluruh workspace (.zip)"),
], 15)

h(doc, "5.5 Auto-save Notebook", 2)
para(doc, "Kode yang diketik pada editor tersimpan otomatis ke /persist (_autosave/paste.ipynb).")
add_shots_grid(doc, [
    ("autosave", "typed", "Mengetik kode pada editor Monaco"),
    ("autosave", "saved-indicator", "Indikator 'tersimpan' otomatis muncul"),
], 15)

h(doc, "5.6 Manajemen Pengguna (Admin)", 2)
para(doc, "Tabel pengguna, pencarian, form tambah pengguna, modal Kelola Kebijakan, dan "
          "menu Aksi (perbaikan BUG-001).")
add_shots_grid(doc, [
    ("users", "list", "Daftar Pengguna"),
    ("users", "search-nomatch", "Pencarian pengguna (contoh: tanpa hasil cocok)"),
    ("users", "add-form", "Form Tambah Pengguna"),
    ("users", "modal-open", "Modal Kelola Kebijakan (read-only saat uji)"),
    ("users", "action-menu", "Menu Aksi per baris (BUG-001 telah diperbaiki)"),
], 15)

h(doc, "5.7 Laporan (Admin)", 2)
para(doc, "Laporan platform: info sistem, penggunaan GPU langsung, pemakaian disk per "
          "user, dan unduh laporan HTML.")
add_shots_grid(doc, [
    ("report", "overview", "Laporan platform (ringkasan)"),
    ("report", "disk-section", "Seksi Pemakaian Disk per User"),
    ("report", "after-download", "Setelah unduh laporan HTML"),
], 15)

h(doc, "5.8 Monitor, Peringatan & Profil", 2)
add_shots_grid(doc, [
    ("monitor", "view", "Monitor sistem — grafik CPU/RAM & per-GPU (admin)"),
    ("alerts", "view", "Halaman Peringatan (admin)"),
    ("alerts", "config-form", "Konfigurasi ambang peringatan"),
    ("profile", "view", "Halaman Profil pengguna"),
], 15)

# ----- 6. PER-PERAN
h(doc, "6. Pengujian Per-Peran (Super Admin · Admin · Dosen · Mahasiswa)", 1)
para(doc, "Otorisasi keempat peran diverifikasi pada tingkat API dan UI. Untuk sisi dosen "
          "dibuat akun QA khusus (CHqadosen) karena sebelumnya belum ada akun dosen.")
h(doc, "6.1 Matriks Otorisasi API", 2)
table(
    doc,
    ["Kasus", "Peran", "/auth/me", "/admin/report", "/monitoring/overview", "Status"],
    [
        ["TC-ROLE-API-superadmin", "Super Admin", "role=admin, is_superadmin=true", "200", "200", "LULUS / SKIP*"],
        ["TC-ROLE-API-admin", "Admin", "role=admin", "200", "200", "LULUS"],
        ["TC-ROLE-API-dosen", "Dosen", "role=dosen", "403 (ditolak)", "200", "LULUS"],
        ["TC-ROLE-API-mahasiswa", "Mahasiswa", "role=mahasiswa", "403 (ditolak)", "200", "LULUS"],
    ],
    widths=[4.2, 2.4, 3.4, 2.6, 2.6, 2.2],
)
para(doc, "* Token QA Super Admin hanya sah bila akun Super Admin sedang memiliki sesi aktif "
          "(kebijakan sesi-tunggal). Bila tidak, kasus ini di-SKIP secara sah agar tidak "
          "mengganggu akun Super Admin yang dipakai pengguna di peramban. Kapabilitas "
          "'require_admin' tetap terbukti melalui peran Admin (peran identik).",
     size=9, italic=True, color=MUTED)
h(doc, "6.2 UI Role-Aware — Mahasiswa", 2)
add_shots_grid(doc, [
    ("roles-mahasiswa", "dashboard", "Dashboard sisi Mahasiswa — judul 'Ruang Belajar Mahasiswa', sidebar tanpa menu admin"),
    ("roles-mahasiswa", "users-blocked", "Mahasiswa membuka /users — tidak ada data pengguna yang bocor"),
], 15)
h(doc, "6.3 UI Role-Aware — Dosen", 2)
add_shots_grid(doc, [
    ("roles-dosen", "dashboard", "Dashboard sisi Dosen — judul 'Ruang Kerja Dosen', sidebar tanpa menu admin"),
    ("roles-dosen", "users-blocked", "Dosen membuka /users — tidak ada data pengguna yang bocor"),
], 15)
para(doc, "Menu yang tersembunyi dari mahasiswa/dosen: Monitor, Laporan, Peringatan, "
          "Pengguna, Pengaturan.")

# ----- 7. KEAMANAN
h(doc, "7. Pengujian Keamanan", 1)
table(
    doc,
    ["ID", "Aspek", "Hasil"],
    [
        ["SEC-01", "Header keamanan & anti-clickjacking (X-Frame-Options, CSP, HSTS, nosniff)", "LULUS"],
        ["SEC-02", "Endpoint terproteksi menolak tanpa token (401)", "LULUS"],
        ["SEC-03", "Privilege escalation: mahasiswa → endpoint admin ditolak (403)", "LULUS"],
        ["SEC-04", "Directory traversal pada workspace ditolak", "LULUS"],
        ["SEC-05", "SQL injection pada login ditolak", "LULUS"],
        ["SEC-06", "CORS tidak memantulkan origin jahat", "LULUS"],
        ["SEC-07", "Tidak ada kebocoran stack trace / path internal (404)", "LULUS"],
        ["SEC-08", "XSS pada kolom cari tidak tereksekusi", "LULUS"],
        ["SEC-09", "Token tidak bocor ke cookie / URL", "LULUS"],
    ],
    widths=[2, 11, 2.5],
)
add_shot(doc, "security", "xss-probe", "Uji XSS pada kolom pencarian — payload tidak tereksekusi (tidak ada dialog)")

# ----- 8. PERFORMA
h(doc, "8. Pengujian Performa", 1)
para(doc, "Latensi endpoint diukur setelah migrasi basis data ke PostgreSQL lokal.")
table(
    doc,
    ["Metrik", "Ambang", "Hasil", "Status"],
    [
        ["GET /health", "< 1500 ms", "± 16 ms", "LULUS"],
        ["GET /admin/report", "< 8000 ms", "± 30 ms", "LULUS"],
        ["Muat halaman (×5 rute)", "wajar", "cepat", "LULUS"],
        ["Konkurensi ringan", "stabil", "stabil", "LULUS"],
    ],
    widths=[5, 3.5, 3.5, 3],
)

# ----- 9. RESPONSIF
h(doc, "9. Pengujian Responsif (Mobile & Tablet)", 1)
para(doc, "Tata letak diuji pada mobile (412×839) dan tablet (820×1180): memuat tanpa "
          "error dan tanpa overflow horizontal berlebih.")
h(doc, "9.1 Mobile (412×839)", 2)
add_shots_grid(doc, [
    ("responsive/412x839", "root", "Mobile — Dashboard"),
    ("responsive/412x839", "_storage", "Mobile — Penyimpanan (BUG-002 diperbaiki: tanpa overflow)"),
    ("responsive/412x839", "_jobs", "Mobile — Daftar Job"),
    ("responsive/412x839", "_report", "Mobile — Laporan"),
], 7.5)
h(doc, "9.2 Tablet (820×1180)", 2)
add_shots_grid(doc, [
    ("responsive/820x1180", "root", "Tablet — Dashboard"),
    ("responsive/820x1180", "_storage", "Tablet — Penyimpanan"),
], 11)

# ----- 10. BUG
h(doc, "10. Temuan (Bug) & Perbaikan", 1)
h(doc, "BUG-001 — Menu 'Aksi' menutup sendiri saat tabel overflow (UX)", 2)
table(doc, ["Field", "Nilai"], [
    ["Severity / Priority", "Low–Medium / Medium"],
    ["Komponen", "frontend/src/pages/Users.tsx → RowActions"],
    ["Gejala", "Menu Aksi tertutup seketika saat tabel Pengguna overflow horizontal (focus-scroll)."],
    ["Perbaikan", "Menu di-render via portal + position:fixed + abaikan scroll 350ms pertama saat buka."],
    ["Status", "DIPERBAIKI — dikunci TC-USR-03 (3/3 lulus)."],
], widths=[4, 12])
h(doc, "BUG-002 — Header Penyimpanan overflow horizontal di mobile (UX responsif)", 2)
table(doc, ["Field", "Nilai"], [
    ["Severity / Priority", "Low / Medium"],
    ["Komponen", "frontend/src/pages/Storage.tsx → header"],
    ["Ditemukan oleh", "TC-RESP /storage (mobile 412×839) saat menambah fitur unduh folder"],
    ["Gejala", "Tombol baru 'Unduh semua' membuat baris header melebihi lebar layar (overflow 64px)."],
    ["Perbaikan", "Baris tombol dibuat membungkus: flex flex-wrap items-center justify-end gap-2 sm:gap-3."],
    ["Status", "DIPERBAIKI — dikunci TC-RESP /storage (mobile & tablet)."],
], widths=[4, 12])
para(doc, "Kedua bug bersifat kosmetik/UX (bukan cacat fungsional atau keamanan) dan telah "
          "ditutup pada sesi yang sama; regresinya dikunci oleh kasus uji.", italic=True, color=MUTED)

# ----- 11. OBSERVASI
h(doc, "11. Observasi (bukan cacat)", 1)
table(
    doc,
    ["ID", "Tipe", "Catatan"],
    [
        ["OBS-1", "Performa", "TERATASI — migrasi ke PostgreSQL lokal menurunkan latensi endpoint drastis."],
        ["OBS-2", "Desain", "/jobs hanya filter status (tanpa pencarian teks) — by design."],
        ["OBS-3", "Lingkungan", "Mode headed tak tersedia (server headless); diverifikasi via video+trace+screenshot."],
        ["OBS-4", "Keamanan", "Token di localStorage (wajar utk SPA bearer; dimitigasi CSP; HSTS terpasang)."],
        ["OBS-5", "Ketahanan", "Editor Monaco dimuat dari CDN jsdelivr; saran: bundel same-origin agar tahan jaringan lambat/offline."],
        ["OBS-6", "Uji peran", "Token QA Super Admin hanya sah saat ada sesi aktif; bila tidak → skip sah (sesi-tunggal)."],
    ],
    widths=[1.6, 2.4, 12],
)

# ----- 12. INVENTARIS
h(doc, "12. Inventaris Kasus Uji", 1)
table(
    doc,
    ["Area", "Kasus uji", "Hasil"],
    [
        ["Publik/Auth", "Landing, Login, validasi, login salah, toggle password, 404, redirect (7)", "7 LULUS"],
        ["API", "health, auth, schema, authz, disk, payload, 404, latensi, report/user (10)", "10 LULUS"],
        ["API-Peran", "matriks 4 peran (super/admin/dosen/mahasiswa)", "3 LULUS · 1 SKIP*"],
        ["Keamanan", "SEC-01..SEC-09", "9 LULUS"],
        ["Dashboard/Navigasi", "Dashboard, semua rute, sidebar, back/forward/refresh", "LULUS"],
        ["Daftar Job", "list/empty, filter, detail (search N/A)", "LULUS (1 skip sah)"],
        ["Penyimpanan", "kuota, unggah+bersih, tombol unggah, unduh workspace/folder .zip (TC-STO-01..06)", "6 LULUS"],
        ["Auto-save", "ketik → tersimpan ke _autosave/paste.ipynb", "LULUS"],
        ["Pengguna", "tabel, cari, modal kebijakan, form tambah (TC-USR-01..04)", "LULUS"],
        ["Laporan", "seksi sistem, disk per user, unduh HTML", "LULUS"],
        ["Peran-UI", "mahasiswa & dosen: dashboard + sidebar + /users (TC-ROLE-*)", "4 LULUS"],
        ["Responsif", "mobile & tablet: /, /jobs, /storage, /report + navigasi", "LULUS"],
        ["Performa", "latensi + muat halaman + konkurensi ringan", "LULUS"],
    ],
    widths=[3.2, 10.5, 2.5],
)
para(doc, "* Super Admin di-skip sah bila tak ada sesi aktif (lihat OBS-6).", size=9, italic=True, color=MUTED)

# ----- 13. KESIMPULAN
h(doc, "13. Kesimpulan", 1)
para(
    doc,
    "Aplikasi UNISMUH ComputeHub LULUS pengujian menyeluruh: 90 dari 92 kasus lulus, 2 "
    "di-skip secara sah, tanpa kegagalan fungsional, tanpa celah otorisasi/injeksi, dan "
    "tanpa error fatal. Fitur baru unduh folder/workspace (.zip) berfungsi benar pada API "
    "maupun UI. Otorisasi keempat peran (Super Admin, Admin, Dosen, Mahasiswa) berperilaku "
    "sesuai rancangan — endpoint admin menolak dosen & mahasiswa (403), dan antarmuka "
    "menyesuaikan peran. Dua temuan UX telah diperbaiki dan dikunci oleh kasus uji.",
)
para(
    doc,
    "Aplikasi dinilai LAYAK untuk digunakan. Rekomendasi peningkatan (tidak menghambat): "
    "membundel editor Monaco secara same-origin untuk ketahanan jaringan kampus (OBS-5).",
    bold=True,
)

# ----- LAMPIRAN
h(doc, "Lampiran: Cara Menjalankan & Artefak", 1)
para(doc, "Menjalankan ulang seluruh suite:")
para(doc, "cd testing && npx playwright test", size=10, italic=True)
para(doc, "Laporan interaktif (langkah + video + trace):")
para(doc, "npm run report   →   reports/html-report/index.html", size=10, italic=True)
bullets(doc, [
    "reports/junit/results.xml — hasil untuk CI",
    "reports/json/results.json — hasil mesin",
    "screenshots/ · test-results/ (video, trace) — artefak visual per langkah",
    "bug-report.md · security-report.md · performance-report.md · coverage-report.md · blackbox/*.md",
])

doc.save(str(OUT))
print("SAVED", OUT)
print("SIZE", OUT.stat().st_size, "bytes")
